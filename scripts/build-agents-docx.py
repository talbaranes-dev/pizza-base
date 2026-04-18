"""
Build מערך-סוכנים-להקמת-אתרים.docx from the current .claude/agents/*.md files.

Usage:
    python scripts/build-agents-docx.py

Run any time you edit an agent md file. Keeps the docx in lockstep with the specs.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
AGENTS_DIR = ROOT / ".claude" / "agents"
OUT = ROOT / "מערך-סוכנים-להקמת-אתרים.docx"

ORDER = [
    ("orchestrator", "המתזמר — נקודת הכניסה"),
    ("template-agent", "שכפול תבנית — פיצה base → פיצרייה חדשה"),
    ("firebase-agent", "Firebase — יצירת פרויקט חדש"),
    ("domain-agent", "בדיקת בעלות על הדומיין"),
    ("dns-agent", "ניהול רשומות DNS ב-JetDNS"),
    ("deploy-agent", "Build ו-Deploy"),
    ("verify-agent", "אימות שהאתר באוויר"),
]


def set_rtl(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    return h


def add_para(doc, text, *, bold=False, mono=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def parse_agent(md_path: Path):
    content = md_path.read_text(encoding="utf-8")
    parts = content.split("---", 2)
    assert len(parts) >= 3, f"bad frontmatter in {md_path}"
    frontmatter, body = parts[1], parts[2].strip()
    meta = {}
    for line in frontmatter.strip().splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip()] = v.strip()
    return meta, body


def main():
    doc = Document()

    title = doc.add_heading("מערך סוכנים להקמת פיצרייה חדשה", level=0)
    set_rtl(title)

    add_para(
        doc,
        "מטרת המערכת: לקבל שם של פיצרייה חדשה, ולהפיק אתר חי בכתובת "
        "https://<שם>.bybe.co.il/ — מקצה לקצה, ללא התערבות ידנית. "
        "התבנית ב-pizza-base-main משתמשת ב-placeholders טהורים (YOUR_PROJECT_ID, YOUR_DISPLAY_NAME) "
        "והסוכנים ממלאים אותם. אתרי דמו חיים להשוואה ויזואלית: "
        "pizzademoorder.bybe.co.il ו-pizzademoadmin.bybe.co.il.",
    )

    add_heading(doc, "שרשרת ההפעלה", level=1)
    add_para(
        doc,
        "המשתמש אומר ״הקם פיצרייה <שם>״. Orchestrator מפעיל את הסוכנים בסדר הזה:",
    )
    stages = [
        "template-agent — שכפול פיצה base לתיקייה חדשה ב-BYBE\\<שם>\\, והחלפת הפלייסהולדרים (YOUR_PROJECT_ID, YOUR_DISPLAY_NAME) בערכים החדשים.",
        "firebase-agent — יצירת פרויקט Firebase חדש, שני Hosting sites, Realtime Database, משתמש auth.",
        "template-agent (שלב שני) — הזרקת API key ו-databaseURL החדשים אל הקבצים המשוכפלים.",
        "domain-agent — אימות שאנחנו שולטים ב-bybe.co.il (nameservers ב-JetDNS).",
        "firebase-agent — הוספת <שם>.bybe.co.il ל-Hosting, קבלת TXT לאימות.",
        "dns-agent — כתיבת ה-TXT, המתנה להתפשטות (עד 10 דקות).",
        "firebase-agent — הפעלת אימות הדומיין, קבלת רשומות A הנדרשות.",
        "dns-agent — כתיבת ה-A records, המתנה להתפשטות.",
        "deploy-agent — firebase deploy מהתיקיות המשוכפלות.",
        "verify-agent — HTTP 200 ו-SSL תקין על https://<שם>.bybe.co.il.",
    ]
    for i, s in enumerate(stages, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        set_rtl(p)

    add_heading(doc, "שמירת מצב ו-Rollback", level=2)
    add_para(
        doc,
        "אחרי כל שלב ה-Orchestrator כותב ל-<target_root>\\.orchestrator-state.json — "
        "כך אפשר לחזור על ריצה חלקית בלי להתחיל מחדש. "
        "Rollback ברירת מחדל מחזיר רק רשומות DNS ו-Hosting releases. "
        "מחיקת פרויקט Firebase דורשת אישור מפורש מהמשתמש.",
    )

    # Per-agent sections
    for slug, heading in ORDER:
        md_path = AGENTS_DIR / f"{slug}.md"
        if not md_path.exists():
            continue
        meta, body = parse_agent(md_path)
        add_heading(doc, heading, level=1)
        add_para(doc, meta.get("description", ""), bold=True)

        # Pull the Procedure section if present, else use the first two paragraphs
        lines = body.splitlines()
        current = []
        capture = False
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("## Procedure") or stripped.startswith("## Operations"):
                capture = True
                continue
            if capture and stripped.startswith("## "):
                break
            if capture:
                current.append(line)
        if not current:
            current = lines[:20]
        snippet = "\n".join(current).strip()
        if len(snippet) > 1500:
            snippet = snippet[:1500] + "\n..."
        add_code(doc, snippet)

    add_heading(doc, "דוגמה מלאה", level=1)
    add_para(
        doc,
        "משתמש: ״הקם פיצרייה pizza-gluten״ → Orchestrator יוצר "
        "BYBE\\pizza-gluten\\, משכפל את פיצה base, בונה פרויקט Firebase חדש, "
        "מוסיף DNS ב-JetDNS, ומדפלוי. תוצאה סופית: "
        "https://pizza-gluten.bybe.co.il/ חי בערך 15 דקות אחרי הפקודה.",
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
